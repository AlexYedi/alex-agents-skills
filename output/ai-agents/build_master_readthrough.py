"""
SUBSTRATE · Vol III · MASTER READ-THROUGH
A docx that provides the learning path for the entire ai-agents module:
  - Auto-updating Word table of contents
  - The Learning Path (sequenced 90-min and full-cycle reads)
  - The Plates (summary of each plate and what to extract)
  - The Readings (summary of Report + Addendum chapters)
  - The Projects / To-Dos (action map, bet next-actions, cruxes to watch)

Output: AI_AGENTS_MASTER_READTHROUGH.docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = "/Users/sameoldexpressions/Documents/GitHub/alex-agents-skills/output/ai-agents"
OUT_PATH = f"{BASE}/AI_AGENTS_MASTER_READTHROUGH.docx"

# ---------- Palette ----------
INK = RGBColor(0x17, 0x12, 0x10)
INK_SOFT = RGBColor(0x3A, 0x2F, 0x26)
GRAY = RGBColor(0x7A, 0x6E, 0x60)
VERMILION = RGBColor(0xA6, 0x37, 0x1F)
OCHRE = RGBColor(0xA3, 0x74, 0x25)
RULE = RGBColor(0x2A, 0x21, 0x1B)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_run_font(run, *, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # Ensure East-Asian fallback uses same font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)


def add_para(doc, text="", *, style=None, font="Calibri", size=11,
             bold=False, italic=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        set_run_font(r, name=font, size=size, bold=bold, italic=italic, color=color)
    return p


def add_horizontal_rule(doc, color=RULE):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    hex_color = "{:02X}{:02X}{:02X}".format(*color)
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_toc_field(doc):
    """Insert a Word TOC field. User must Update Field on open (Word does this on prompt)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # Heading 1 = level 1; Heading 2 = level 2; Heading 3 = level 3; with hyperlinks
    instrText.text = r'TOC \o "1-3" \h \z \u'

    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')

    placeholder = OxmlElement('w:r')
    placeholder_text = OxmlElement('w:t')
    placeholder_text.text = "Right-click and choose Update Field to populate this table of contents."
    placeholder.append(placeholder_text)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._element.append(fldChar_begin)
    run._element.append(instrText)
    run._element.append(fldChar_sep)
    run._element.append(placeholder)
    run._element.append(fldChar_end)


def add_heading(doc, text, level=1):
    """Use built-in heading styles so TOC picks them up."""
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, name="Georgia", size=22, bold=True, color=INK)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        set_run_font(r, name="Georgia", size=15, bold=True, color=INK)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(r, name="Georgia", size=12, bold=True, color=VERMILION)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, *, level=0, font="Calibri", size=11,
               bold=False, italic=False, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, name=font, size=size, bold=bold, italic=italic,
                 color=color if color else INK)
    return p


def add_callout(doc, label, body, color=VERMILION):
    """Two-line callout: an eyebrow label and a body sentence."""
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(label.upper())
    set_run_font(r1, name="Consolas", size=8, bold=True, color=color)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    r2 = p2.add_run(body)
    set_run_font(r2, name="Georgia", size=11, italic=True, color=INK_SOFT)


# ============================================================
# CONTENT BUILDERS
# ============================================================

def build_cover(doc):
    # Eyebrow
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("A N   A T L A S   O F   T H E   A G E N T   S T R A T A")
    set_run_font(r, name="Consolas", size=9, color=GRAY)

    add_horizontal_rule(doc)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("The Agent, Seen Whole")
    set_run_font(r, name="Georgia", size=34, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("Master Read-Through · Learning Path · Project Map")
    set_run_font(r, name="Georgia", size=14, italic=True, color=INK_SOFT)

    add_horizontal_rule(doc, color=VERMILION)

    # Subtitle / description
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        "Volume III of the OCQ corpus, distilled into one guided sequence. "
        "How to read the readings, when to look at the plates, what to do once "
        "you have."
    )
    set_run_font(r, name="Georgia", size=11, italic=True, color=INK_SOFT)

    # Compiled-for block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Compiled for A. Yedi")
    set_run_font(r, name="Georgia", size=12, italic=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Cycle MMXXVI · Revision I · 2026-05-14")
    set_run_font(r, name="Consolas", size=9, color=GRAY)

    add_page_break(doc)


def build_toc(doc):
    add_heading(doc, "Contents", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "Open in Word and choose Update Field (or press F9) to populate this "
        "table of contents from the headings below."
    )
    set_run_font(r, name="Georgia", size=10, italic=True, color=GRAY)
    add_toc_field(doc)
    add_page_break(doc)


def build_how_to_read(doc):
    add_heading(doc, "How to read this volume", level=1)

    add_para(doc,
        "The agent layer is the same six artifacts looked at from six angles. "
        "This volume tells you which angle to start from, what each artifact "
        "is actually for, and what action drops out the other end. Two paths "
        "are offered — pick the one that fits the time you have.",
        font="Georgia", size=11)

    add_heading(doc, "Path A — The 90-minute speed-read", level=2)
    add_para(doc,
        "If you only have an evening, take this route. You will lose the texture "
        "of the descriptive volume but you will keep the operating decisions intact.",
        font="Georgia", size=10.5, italic=True, color=INK_SOFT)
    speed = [
        ("5 min",  "Plate M — Master Plate · the synthesis frontispiece"),
        ("10 min", "Tracker §A — the seven Bets refreshed (Vol III deltas)"),
        ("20 min", "Addendum Part XI — Synthesis · Bet rewrites · Risks · Cruxes"),
        ("15 min", "Addendum Part XIII — the Agent-Specific Procurement Rubric"),
        ("10 min", "Plate I (index) + Plate V (vertical products)"),
        ("10 min", "Report Executive Summary — five forces structuring the layer"),
        ("20 min", "Addendum Part XII — the 6 / 12 / 18-Month Action Map"),
    ]
    for t, label in speed:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{t}  ")
        set_run_font(r1, name="Consolas", size=10, bold=True, color=OCHRE)
        r2 = p.add_run(label)
        set_run_font(r2, name="Georgia", size=11, color=INK)

    add_heading(doc, "Path B — The full cycle (≈ 6 hours)", level=2)
    add_para(doc,
        "The reading order in which every later piece assumes the earlier one. "
        "Each step has a 'so what' — the single takeaway you should walk away with.",
        font="Georgia", size=10.5, italic=True, color=INK_SOFT)
    full = [
        ("Step 1", "Plate M — Master Plate",
         "The map of the territory. Every later plate is a zoom into one quadrant."),
        ("Step 2", "Report (AI_AGENTS_REPORT.docx) — Foundation Reference",
         "The descriptive layer. Read in four parts: Capability Substrate · Loops of Cognition · Action & Defense · Productized Layer."),
        ("Step 3", "Plates I–VI — Substrate Vol I · the agent column",
         "Visual atlas of the ten sub-strata plus four meta-strata. Each plate is the Report's same content as a single page you can hold in mind."),
        ("Step 4", "Addendum (AI_AGENTS_ADDENDUM.docx) — Decisions Playbook",
         "Five frameworks applied to the strata. Read Parts VI–X for analysis, then XI–XIII for the operating decisions."),
        ("Step 5", "Plates VII–XI — Substrate Vol II · framework plates",
         "Same five frameworks as Plates: OCQ heat · Wardley · 7 Powers · JTBD · Action Portfolio."),
        ("Step 6", "Tracker (AI_AGENTS_TRACKER.md) — Living Tracker",
         "The working surface. Three sections, three cadences: monthly Bets, bi-weekly talent + capital, trigger-based cruxes and risks."),
    ]
    for step, label, takeaway in full:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(6)
        r1 = p.add_run(f"{step}  ")
        set_run_font(r1, name="Consolas", size=9, bold=True, color=VERMILION)
        r2 = p.add_run(label)
        set_run_font(r2, name="Georgia", size=11.5, bold=True, color=INK)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(8)
        p2.paragraph_format.left_indent = Inches(0.35)
        r3 = p2.add_run("So what  ·  ")
        set_run_font(r3, name="Consolas", size=8.5, color=OCHRE)
        r4 = p2.add_run(takeaway)
        set_run_font(r4, name="Georgia", size=10.5, italic=True, color=INK_SOFT)

    add_callout(doc, "The Vol III delta in one line",
        "Bet #1 is sequenced first (claim Process Power via the Procurement "
        "Playbook). Bet #2 second (collect equity at a vertical-agent winner). "
        "Bet #3 third (compound as advisory). Bets #4 and #5 fold into #1 as modules.")

    add_page_break(doc)


def build_the_plates(doc):
    add_heading(doc, "The plates", level=1)
    add_para(doc,
        "Twelve plates in the order they should be experienced. The visual "
        "atlas lives in AI_AGENTS_MASTER_PLATES.pdf — this section tells you "
        "what each plate is for and what to extract from it.",
        font="Georgia", size=11, italic=True, color=INK_SOFT)

    add_heading(doc, "Plate M — Master Plate", level=2)
    add_para(doc,
        "The synthesis frontispiece. The agent layer drawn as one connected "
        "system: ten sub-strata stacked vertically, four meta-strata wrapping "
        "horizontally, the seven Bets sketched as actions in the margins.",
        font="Georgia", size=11)
    add_callout(doc, "Extract",
        "Where on this single page do you most want to put your next 12 months "
        "of attention? That quadrant is your reading lens for everything that follows.")

    add_heading(doc, "Substrate Volume I — Plates I–VI · the agent column", level=2)
    column_plates = [
        ("Plate I",   "Index of the agent column",
         "The ten sub-strata named and stacked. The reading order for the "
         "Report. Every other plate is a zoom into one of these rows."),
        ("Plate II",  "The cognition loop",
         "Perceive · plan · act · reflect. Read this before Report Part II so "
         "the planning / reasoning vocabulary lands precisely."),
        ("Plate III", "Memory and state",
         "Short-term scratchpad, episodic, long-term, shared. The single "
         "plate that separates demo-grade agents from production-grade."),
        ("Plate IV",  "Action surfaces",
         "Tool use, browser, code interpreter, computer use, voice. Each "
         "surface has its own safety story — read with Report Part III open."),
        ("Plate V",   "Vertical products",
         "Where the agent meets the buyer. Sierra · Decagon · Glean · "
         "Harvey · Hippocratic. This is Bet #2's hunting ground."),
        ("Plate VI",  "Evaluation · observability · safety · regulation",
         "The four meta-strata that wrap the column. The plate that makes "
         "Bet #1's procurement angle visible."),
    ]
    for label, title, body in column_plates:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(6)
        r1 = p.add_run(f"{label}  ·  ")
        set_run_font(r1, name="Consolas", size=9.5, bold=True, color=OCHRE)
        r2 = p.add_run(title)
        set_run_font(r2, name="Georgia", size=11.5, bold=True, color=INK)
        add_para(doc, body, font="Georgia", size=10.5, color=INK_SOFT, space_after=4)

    add_heading(doc, "Substrate Volume II — Plates VII–XI · the framework plates", level=2)
    framework_plates = [
        ("Plate VII",  "OCQ heat map",
         "Opportunities · Challenges · Open Questions painted across every "
         "stratum. The plate that names where the gaps are."),
        ("Plate VIII", "Wardley map of the agent stack",
         "Genesis → Custom → Product → Commodity for every component. The "
         "claimable flags. Where Bet #3 lives."),
        ("Plate IX",   "Helmer's 7 Powers at the agent layer",
         "Which of the seven powers actually compound here. Procurement "
         "Process Power at the meta-layer is Alex's only durable claim."),
        ("Plate X",    "Six ecosystem-level jobs to be done",
         "The Job 6 phase map IS the Procurement Playbook's table of contents. "
         "Read this before drafting Bet #1."),
        ("Plate XI",   "The Action Portfolio",
         "The seven Bets refreshed, sequenced, sized. Compare to Tracker §A "
         "for the same content in update-cadence form."),
    ]
    for label, title, body in framework_plates:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(6)
        r1 = p.add_run(f"{label}  ·  ")
        set_run_font(r1, name="Consolas", size=9.5, bold=True, color=OCHRE)
        r2 = p.add_run(title)
        set_run_font(r2, name="Georgia", size=11.5, bold=True, color=INK)
        add_para(doc, body, font="Georgia", size=10.5, color=INK_SOFT, space_after=4)

    add_page_break(doc)


def build_the_readings(doc):
    add_heading(doc, "The readings", level=1)
    add_para(doc,
        "Three prose artifacts make up the descriptive and decisive layers. "
        "Read them in the order below; each later document assumes the earlier one.",
        font="Georgia", size=11, italic=True, color=INK_SOFT)

    # --- Report ---
    add_heading(doc, "AI_AGENTS_REPORT.docx — Foundation Reference (~10,000 words)", level=2)
    add_para(doc,
        "The descriptive volume. Walks every stratum top-to-bottom in prose. "
        "Read with Plates I–VI open beside you — the plates are the same content "
        "drawn for memory; the report is the same content drawn for argument.",
        font="Georgia", size=11)
    add_para(doc, "Contents in reading order:", font="Georgia", size=10.5,
             italic=True, color=INK_SOFT)
    report_parts = [
        ("Preamble + Executive Summary",
         "The five forces structuring the agent layer in one breath. "
         "Read first, then again last."),
        ("Part I · The Capability Substrate",
         "Foundation models with agentic capability · runtimes and harnesses · "
         "the MCP tool-use protocol. Plates I–II."),
        ("Part II · The Loops of Cognition",
         "Memory and state · planning and reasoning · capability and safety. "
         "Plates III + parts of VI."),
        ("Part III · Action and Defense Surfaces",
         "Action surfaces · evaluation, observability, safety · regulation. "
         "Plates IV + VI."),
        ("Part IV · The Productized Layer",
         "Vertical products · end-user surfaces. Plate V. The bridge from "
         "the descriptive layer into Bet #2's hunting ground."),
    ]
    for title, body in report_parts:
        add_bullet(doc, title, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.55)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(body)
        set_run_font(r, name="Georgia", size=10.5, italic=True, color=INK_SOFT)

    # --- Addendum ---
    add_heading(doc, "AI_AGENTS_ADDENDUM.docx — Decisions Playbook (~17,000 words)", level=2)
    add_para(doc,
        "The decisive volume. Five frameworks applied at agent-layer resolution, "
        "then collapsed into the seven Bets, five Risks, five Cruxes, and the "
        "six-month action map. Read with Plates VII–XI open beside you.",
        font="Georgia", size=11)
    add_para(doc, "Contents in reading order:", font="Georgia", size=10.5,
             italic=True, color=INK_SOFT)
    addendum_parts = [
        ("Part VI · Methodology",
         "How the five lenses inherit from and depart from the parent OCQ tracker."),
        ("Part VII · The OCQ × Agent Sub-Stratum Matrix",
         "Opportunities, challenges, open questions per stratum. Plate VII."),
        ("Part VIII · Wardley Mapping the Agent Stack",
         "Component evolution map and the claimable flags. Plate VIII."),
        ("Part IX · Helmer's 7 Powers and Agent-Layer Durability",
         "Which of the seven powers actually hold at this layer. Plate IX."),
        ("Part X · Ecosystem-Level Jobs to be Done",
         "The six jobs the agent layer performs for buyers. Plate X."),
        ("Part XI · Synthesis — Seven Bets Refreshed, Five Risks, Five Cruxes",
         "The synthesis document. Plate XI. The single chapter to read if you "
         "only have twenty minutes."),
        ("Part XII · 6 / 12 / 18-Month Action Map",
         "The operating playbook. The Projects section of this volume is the "
         "rendered form of this part."),
        ("Part XIII · The Agent-Specific Procurement Rubric",
         "The Appendix to Bet #1. The eight-section rubric that is the "
         "Procurement Playbook's nucleus."),
    ]
    for title, body in addendum_parts:
        add_bullet(doc, title, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.55)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(body)
        set_run_font(r, name="Georgia", size=10.5, italic=True, color=INK_SOFT)

    # --- Tracker ---
    add_heading(doc, "AI_AGENTS_TRACKER.md — Living Tracker", level=2)
    add_para(doc,
        "The working surface. Not a document to be read once and shelved — "
        "the live instrument of attention. Three sections at three cadences. "
        "Update on schedule; the tracker becomes useless the moment it goes stale.",
        font="Georgia", size=11)
    tracker_sections = [
        ("§A · The seven Bets (agent-layer refresh)",
         "Update monthly: conviction, leading indicators, dead modules."),
        ("§B · Agent-specific talent + capital flow",
         "Update bi-weekly: senior moves, $50M+ rounds, ARR watchlist, "
         "NYC peak hiring snapshot."),
        ("§C · Agent-layer cruxes and decision triggers",
         "Update when triggered: each crux has a measurable answer-event."),
        ("§D · Agent-layer structural risks",
         "Update when triggered: each risk has a watch-signal."),
        ("§E · Update log",
         "Append-only. Every change to §A–§D logged with date and one-line why."),
    ]
    for title, body in tracker_sections:
        add_bullet(doc, title, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.55)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(body)
        set_run_font(r, name="Georgia", size=10.5, italic=True, color=INK_SOFT)

    add_page_break(doc)


def build_projects(doc):
    add_heading(doc, "Projects and to-dos", level=1)
    add_para(doc,
        "What drops out of the volume. Drawn from Addendum Part XII (the action "
        "map), Part XIII (the procurement rubric), and Tracker §A's per-Bet "
        "next-actions. Sequenced over an eighteen-month horizon.",
        font="Georgia", size=11, italic=True, color=INK_SOFT)

    # ---- Months 0–6 ----
    add_heading(doc, "Months 0–6 · Plant the procurement flag, run the role search in parallel", level=2)
    add_callout(doc, "Operating Posture",
        "Bet #1 in build mode, Bet #2 in scan mode, Bet #6 in publish mode, "
        "Bet #4 as a single audit pilot folded into Bet #1.")
    m06 = [
        ("Bet #1 · Procurement Playbook outline (Week 1)",
         "Use the Job 6 phase map and the A6 five-questions structure as the "
         "table of contents and chapter prompts."),
        ("Bet #1 · 30 expert interviews (Weeks 2–6)",
         "CISOs · procurement leads · GRC heads at AI-adopting enterprises. "
         "Confirm the eight-section rubric stress-tests against real buying "
         "behavior."),
        ("Bet #1 · Playbook publish (Week 12)",
         "Open download of the Agent Procurement Rubric attached. "
         "Falsifiability test = 500 downloads / 50 inbound by Month 6."),
        ("Bet #2 · NYC role search open (Week 1)",
         "Target list of ten: Sierra (highest-aligned) · Decagon · Glean · "
         "Harvey · Hippocratic · Hebbia · Rogo · Clay · Runway · Ramp AI org. "
         "Comp band $250–400K base + 0.05–0.40% equity."),
        ("Bet #6 · Newsletter Kit v1 publish (Week 4)",
         "Decide weekly vs. biweekly cadence by Week 8. Cross-pollinate "
         "Bet #1 Playbook downloads with subscribers — correlation = "
         "right-audience proof."),
        ("Bet #4 · Trajectory cost-and-latency audit pilot",
         "Free first audit for one mid-market AI-spending company. Case "
         "study by Month 6. Fold into Bet #1 as a module."),
        ("Crux watch · Anthropic ARR resolution Q3 2026",
         "Single load-bearing Bet #2 timing variable. Sign before resolution "
         "if conviction high; after if wobbly."),
    ]
    for title, body in m06:
        add_bullet(doc, title, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.55)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(body)
        set_run_font(r, name="Georgia", size=10.5, italic=True, color=INK_SOFT)

    # ---- Months 6–12 ----
    add_heading(doc, "Months 6–12 · Commit to whichever survived", level=2)
    add_callout(doc, "Operating Posture",
        "Falsifiability tests resolve. Survivors commit, casualties killed. "
        "Bet #3 audit phase begins.")
    m612 = [
        ("Bet #1 · Advisory phase commit (if Playbook clears the bar)",
         "3–5 paid engagements at $25–100K each. At least two Procurement "
         "Audit reports and one architecture review."),
        ("Bet #2 · Tier-2 expansion or kill (if no offer by Month 6)",
         "Expand to Decagon NYC, Regal, smaller bands. No offer by Month 9 "
         "= bet on Alex's profile is wrong OR market is contracting; recheck Bet #7."),
        ("Bet #6 · Cadence commit or kill (5,000 subs and 3+ inbound)",
         "Weekly cadence + paid tier if cleared. Kill at Month 6 per "
         "falsifiability if not."),
        ("Bet #3 · MCP audit of ten SaaS systems (Q3 2026)",
         "The ten enterprise GTM systems most lacking MCP servers, filtered "
         "by Job 4 / Job 1 outcomes. Pair-with-Cloudflare / Kong / Pomerium "
         "positioning. Decide on productization at Month 12."),
        ("Crux watch · MCP commons or fork (EOY 2026)",
         "Major-vendor proprietary schema announcements. A2A adoption "
         "outside Google. Fork = Bet #3 productized form dies; commons = "
         "advisory practice accelerates."),
        ("Crux watch · OSWorld 65% on a frontier system (Q3 2026)",
         "Crossed = computer use moves from supervised to deployable for "
         "narrow back-office a quarter earlier than this volume assumes."),
        ("Crux watch · EU AI Act Article 14 enforcement teeth (Late 2026)",
         "First GPAI fine. Teeth = Bet #1 becomes $10B+ category; paper = "
         "niche but defensible."),
    ]
    for title, body in m612:
        add_bullet(doc, title, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.55)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(body)
        set_run_font(r, name="Georgia", size=10.5, italic=True, color=INK_SOFT)

    # ---- Months 12–18 ----
    add_heading(doc, "Months 12–18 · Compound around the survivors", level=2)
    add_callout(doc, "Operating Posture",
        "Bets that cleared their falsifiability tests now compound. Optional "
        "SaaS productization decision on Bet #1. Bet #7 only as fallback.")
    m1218 = [
        ("Bet #1 · SaaS productization decision (GO / NO)",
         "GO = capitalize. NO = sustain the advisory practice at $250–500K "
         "annual revenue with three-product bundling (Playbook + Architecture "
         "Audit + FinOps for Trajectories)."),
        ("Bet #2 · Deliver-on-Confirm + Modify + Conclude",
         "If a vertical-agent role landed, the first twelve months are about "
         "shipping for that vertical's primary job. The advisory practice "
         "continues at a controlled pace."),
        ("Bet #3 · Productize at 12 or 18 if scar tissue justifies",
         "Otherwise sustain as advisory."),
        ("Bet #7 · Activate only if 1, 2, 3 stall",
         "Build relationships at three NYC funds via RAAIS, Betaworks, "
         "FirstMark MAD. Active applications Q4 2026."),
        ("Cadence · Per-quarter check-in",
         "Hypothesis still holding? Leading indicators moving the right way? "
         "Conviction up, down, or flat? Which bets re-rank from resolved cruxes? "
         "Which bets re-price from materialized risks?"),
    ]
    for title, body in m1218:
        add_bullet(doc, title, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.55)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(body)
        set_run_font(r, name="Georgia", size=10.5, italic=True, color=INK_SOFT)

    # ---- The Procurement Rubric as a separate project ----
    add_heading(doc, "The Procurement Rubric — Bet #1's nucleus", level=2)
    add_para(doc,
        "Addendum Part XIII is both a section of the Playbook and a "
        "standalone downloadable. Eight sections, scored 1–5; below-3 average "
        "fails procurement. The to-do is to harden it against the 30 expert "
        "interviews, then publish it as the open lead-magnet.",
        font="Georgia", size=11)
    rubric = [
        ("1. Define · Map the agent's autonomy scope",
         "What the agent does unattended · what requires human approval · "
         "where decision authority stops."),
        ("2. Locate · Inventory data and tool reach",
         "OAuth blast radius · where memory lives physically · who can subpoena it."),
        ("3. Prepare · Tool-boundary policy and runtime safety stack",
         "Which tools fire unsupervised · which require approval · which are "
         "blocked when context is tainted."),
        ("4. Confirm · Pre-flight evaluation evidence",
         "Signed reproducible eval report with model pin, dataset hash, "
         "harness version. Trajectory regression set sized and curated."),
        ("5. Execute · Action-rollback story",
         "Write tools wrapped in reversible-transaction abstraction. "
         "Documented recovery path for every write."),
        ("6. Monitor · Trajectory observability and silent-failure detection",
         "Detect a trajectory that completes-as-judged but writes the wrong "
         "value. Silent-failure rate disclosed."),
        ("7. Modify · Non-engineer intervention capability",
         "Cockpit UI for line manager redirection. Policy correction shipped "
         "without an engineering ticket."),
        ("8. Conclude · Counterparty evidence and audit pack",
         "EU AI Act / NIST AI RMF / ISO 42001 tie-out · evidence kept fresh "
         "through annual review."),
    ]
    for title, body in rubric:
        add_bullet(doc, title, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.55)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(body)
        set_run_font(r, name="Georgia", size=10.5, italic=True, color=INK_SOFT)

    add_callout(doc, "The closing brief",
        "The procurement-readiness gap is not that buyers don't know what to "
        "ask; it is that they have no signed, falsifiable rubric to ask it "
        "with. Bet #1 is the rubric.")

    add_page_break(doc)


def build_closing(doc):
    add_heading(doc, "Closing — the volume in one breath", level=1)
    add_para(doc,
        "Twelve plates, three readings, one tracker. Sequenced so synthesis "
        "comes first, descriptive layer second, decisive layer third, and "
        "the working surface remains open through the next eighteen months.",
        font="Georgia", size=11)
    add_para(doc,
        "The Vol III thesis is that the agent layer's durable claimable "
        "position for an operator without a research lab is procurement "
        "Process Power at the meta-layer — and that the window to claim it "
        "closes within twelve months. Everything else in the volume is in "
        "service of testing that thesis, sequencing the bets around it, or "
        "instrumenting the fallbacks.",
        font="Georgia", size=11)
    add_para(doc,
        "Read it once. Then update the tracker.",
        font="Georgia", size=11, italic=True, color=INK_SOFT)


# ============================================================
# MAIN
# ============================================================

def main():
    doc = Document()

    # Page setup — Letter, narrow-ish margins
    for section in doc.sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    build_cover(doc)
    build_toc(doc)
    build_how_to_read(doc)
    build_the_plates(doc)
    build_the_readings(doc)
    build_projects(doc)
    build_closing(doc)

    doc.save(OUT_PATH)
    size_kb = os.path.getsize(OUT_PATH) / 1024
    print(f"Wrote {OUT_PATH}")
    print(f"  · {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
